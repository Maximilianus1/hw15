import docx
document1 = docx.Document('HelloWorld.docx')
run1 = document1.paragraphs[0]
for i in run1.runs:
    if i.bold:
        print(i.text)
document2 = docx.Document()
p = document2.add_paragraph('Это интересно! Важно! Молодцы!')
for i in p.runs:
    i.font.size= docx.shared.Pt(30)
    i.font.name = 'Comic Sans MS'
document2.save('test.docx')
